"""Create final report DOCX from markdown content."""
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFENSE_PACKAGE = PROJECT_ROOT / "final_defense_package"


def create_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("ERROR: python-docx not installed.")
        print("Install with: pip install python-docx")
        print("Then run: python scripts/create_final_report_docx.py")
        return False

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('TRƯỜNG ĐẠI HỌC THỦ DẦU MỘT')
    run.bold = True
    run.font.size = Pt(16)

    institute = doc.add_paragraph()
    institute.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = institute.add_run('VIỆN CÔNG NGHỆ SỐ')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    topic = doc.add_paragraph()
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = topic.add_run('ĐỀ TÀI NGHIÊN CỨU KHOA HỌC')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    topic_name = doc.add_paragraph()
    topic_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = topic_name.add_run('Dự đoán nghẽn hệ thống web bằng mô hình trí tuệ nhân tạo dựa trên chuỗi thời gian')
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    doc.add_paragraph()

    # Student info
    students = [
        'Đinh Hữu Phong — 2324802010095 — Nhóm trưởng',
        'Đặng Văn Tuyển — 2324802010156 — Thành viên',
        'Nguyễn Đức Thịnh — 2324802010355 — Thành viên',
    ]
    for student in students:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(student)
        run.font.size = Pt(13)

    doc.add_paragraph()

    advisor = doc.add_paragraph()
    advisor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = advisor.add_run('Giảng viên hướng dẫn: ThS. Nguyễn Ngọc Thận')
    run.font.size = Pt(13)

    doc.add_paragraph()

    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year.add_run('Năm học 2025–2026')
    run.bold = True
    run.font.size = Pt(14)

    # Page break
    doc.add_page_break()

    # Table of Contents
    doc.add_heading('MỤC LỤC', level=1)
    toc_items = [
        'CHƯƠNG 1: GIỚI THIỆU TỔNG QUAN',
        'CHƯƠNG 2: CƠ SỞ LÝ THUYẾT',
        'CHƯƠNG 3: MÔ HÌNH ĐỀ XUẤT VÀ THIẾT KẾ HỆ THỐNG',
        'CHƯƠNG 4: THỰC NGHIỆM VÀ ĐÁNH GIÁ',
        'CHƯƠNG 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # Chapter 1
    doc.add_heading('CHƯƠNG 1: GIỚI THIỆU TỔNG QUAN', level=1)

    doc.add_heading('1.1 Tên đề tài', level=2)
    doc.add_paragraph('Dự đoán nghẽn hệ thống web bằng mô hình trí tuệ nhân tạo dựa trên chuỗi thời gian.')

    doc.add_heading('1.2 Lý do chọn đề tài', level=2)
    doc.add_paragraph(
        'Hệ thống web hiện đại đóng vai trò quan trọng trong đời sống số. '
        'Khi lưu lượng truy cập tăng đột ngột, nguy cơ nghẽn có thể gây mất dữ liệu và thiệt hại kinh tế. '
        'Các phương pháp giám sát truyền thống dựa trên ngưỡng tĩnh có hạn chế lớn. '
        'Trí tuệ nhân tạo có khả năng học pattern phức tạp và dự báo xu hướng tương lai.'
    )

    doc.add_heading('1.3 Mục tiêu', level=2)
    goals = [
        'Xây dựng pipeline dữ liệu chuỗi thời gian từ log web',
        'Thiết kế mô hình TCN-Attention-BiLSTM',
        'Đánh giá với MAE, RMSE, R²',
        'Cảnh báo sớm với threshold calibration',
        'Recommendation Engine đề xuất hành động',
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')

    doc.add_heading('1.4 Giới hạn', level=2)
    doc.add_paragraph(
        'NASA HTTP 1995 là dữ liệu cũ, không có CPU/RAM/response time. '
        'Target là proxy congestion score, không phải measured congestion. '
        'Synthetic stress benchmark chỉ dùng để kiểm thử có kiểm soát.'
    )

    doc.add_page_break()

    # Chapter 2
    doc.add_heading('CHƯƠNG 2: CƠ SỞ LÝ THUYẾT', level=1)

    doc.add_heading('2.1 Hệ thống web và hiệu năng', level=2)
    doc.add_paragraph(
        'Hệ thống web bao gồm web server, application server, database, load balancer. '
        'Các chỉ số hiệu năng: request count, response time, throughput, error rate, CPU/Memory usage.'
    )

    doc.add_heading('2.2 Nghẽn hệ thống web', level=2)
    doc.add_paragraph(
        'Nghẽn xảy ra khi hệ thống không xử lý kịp lưu lượng, dẫn đến tăng response time, tăng error rate, giảm throughput.'
    )

    doc.add_heading('2.3 Các mô hình', level=2)
    models = [
        'LSTM: Long Short-Term Memory, gating mechanism',
        'GRU: Gated Recurrent Unit, đơn giản hơn LSTM',
        'TCN: Temporal Convolutional Network, dilated convolution',
        'Attention: Học trọng số quan trọng',
        'BiLSTM: Bidirectional LSTM, context hai chiều',
    ]
    for model in models:
        doc.add_paragraph(model, style='List Bullet')

    doc.add_page_break()

    # Chapter 3
    doc.add_heading('CHƯƠNG 3: MÔ HÌNH ĐỀ XUẤT', level=1)

    doc.add_heading('3.1 Kiến trúc TCN-Attention-BiLSTM', level=2)
    doc.add_paragraph(
        'Input [batch, 60, 19] → TCN Block → Multi-Head Attention → BiLSTM → Dense → Output [batch, 1]'
    )

    doc.add_heading('3.2 Pipeline dữ liệu', level=2)
    doc.add_paragraph(
        'Raw HTTP Logs → Parse → Aggregate 1-min → Features → Normalize 0-1 → Windows → Split'
    )

    doc.add_heading('3.3 Recommendation Engine', level=2)
    doc.add_paragraph(
        'Risk levels: Normal, Watch, Warning, Critical. '
        'Actions: scale_up_cpu, rate_limit, enable_cache, investigate_anomaly.'
    )

    doc.add_page_break()

    # Chapter 4
    doc.add_heading('CHƯƠNG 4: THỰC NGHIỆM VÀ ĐÁNH GIÁ', level=1)

    doc.add_heading('4.1 Kết quả v2', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    cells[0].text = 'Metric'
    cells[1].text = 'Giá trị'
    data = [
        ('MAE', '0.043053'),
        ('RMSE', '0.056036'),
        ('R²', '0.339994'),
        ('Train time', '1092.5s'),
        ('Threshold', '0.183838'),
        ('Calibrated F1', '0.865596'),
    ]
    for i, (metric, value) in enumerate(data):
        cells = table.rows[i+1].cells
        cells[0].text = metric
        cells[1].text = value

    doc.add_heading('4.2 So sánh mô hình', level=2)
    doc.add_paragraph(
        'TCN-Attention-BiLSTM có R² cao nhất (0.339994) trong 8 mô hình đã thử.'
    )

    doc.add_heading('4.3 Biểu đồ', level=2)
    figures = [
        'outputs/figures/prediction_vs_actual.png',
        'outputs/figures/error_distribution.png',
        'outputs/figures/model_comparison_rmse.png',
        'outputs/figures/training_curves.png',
        'outputs/figures/early_warning_timeline.png',
        'outputs/figures/synthetic_stress_scenarios.png',
    ]
    for fig in figures:
        doc.add_paragraph(f'[Hình: {fig}]', style='List Bullet')

    doc.add_page_break()

    # Chapter 5
    doc.add_heading('CHƯƠNG 5: KẾT LUẬN', level=1)

    doc.add_heading('5.1 Kết quả đạt được', level=2)
    results = [
        'Pipeline dữ liệu chuỗi thời gian',
        'Mô hình TCN-Attention-BiLSTM',
        'MAE 0.043, RMSE 0.056, R² 0.34',
        'Calibrated F1 0.87',
        'Dashboard demo minh bạch',
    ]
    for result in results:
        doc.add_paragraph(result, style='List Bullet')

    doc.add_heading('5.2 Hướng phát triển', level=2)
    future = [
        'Bổ sung dataset mới (Zanbil, Calgary)',
        'Multi-source validation',
        'Tối ưu hyperparameter',
        'Online monitoring',
        'Auto-scaling integration',
    ]
    for item in future:
        doc.add_paragraph(item, style='List Bullet')

    # Save
    output_path = DEFENSE_PACKAGE / "bao_cao_nghien_cuu_final.docx"
    doc.save(str(output_path))
    print(f"OK: Created {output_path}")
    return True


if __name__ == "__main__":
    success = create_docx()
    if not success:
        sys.exit(1)
